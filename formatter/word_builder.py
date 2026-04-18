"""
Word 文件建構模組。

輸入：PaperSchema + ncu.docx 模板路徑
輸出：依學校格式排版完成的 docx.Document 物件

流程：
  1. 開啟模板（繼承所有 Word Styles + 頁邊距 + 頁碼設定）
  2. 保存模板的 embedded sectPr（前置區段：羅馬數字頁碼）
  3. 清空 body 所有段落 / 表格
  4. 寫入前置內容（封面、摘要、誌謝、目錄）
  5. 插入保存的 embedded sectPr（作為前置 / 正文分節點）
  6. 寫入正文各章與參考文獻
  7. 回傳 Document

ncu.docx 模板樣式對照：
  Chapter Heading   → 章標題（置中、粗體、16pt）
  Section Heading   → 節標題（左對齊、粗體、14pt）
  SubSection Heading → 子節標題（左對齊、粗體）
  Body Text         → 內文（兩端對齊）
  Figure Caption    → 圖說（置中、11pt）—— 圖片下方
  Table Caption     → 表說（置中、粗體、11pt）—— 表格上方
  Front Matter Title → 前置區段標題（置中、粗體、16pt）
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from formatter.schema import (
    AbstractSection,
    Chapter,
    ContentBlock,
    CoverInfo,
    FigureBlock,
    PaperSchema,
    ParagraphBlock,
    Section,
    SubSection,
    TableBlock,
)

# ── 中文數字（章序號）────────────────────────────────────────────────────────
_ZH_NUMS = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
            "十一", "十二", "十三", "十四", "十五"]


def _zh_num(n: int) -> str:
    return _ZH_NUMS[n] if n < len(_ZH_NUMS) else str(n)


# ── 段落建構輔助 ──────────────────────────────────────────────────────────────

def _add_para(
    doc: Document,
    text: str,
    style_name: str = "Body Text",
    bold: bool = False,
    font_size_pt: int | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
):
    """新增段落，套用模板樣式；bold / font_size_pt 可覆寫 run 級屬性。"""
    try:
        style = doc.styles[style_name]
    except KeyError:
        style = doc.styles["Normal"]

    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if font_size_pt:
        run.font.size = Pt(font_size_pt)
    if alignment is not None:
        p.alignment = alignment
    return p


def _add_front_matter_title(doc: Document, text: str):
    """前置區段標題（摘要、誌謝、目錄、參考文獻）→ Front Matter Title 樣式。"""
    return _add_para(doc, text, "Front Matter Title")


def _add_chapter_heading(doc: Document, text: str):
    """章標題 → Chapter Heading 樣式。"""
    return _add_para(doc, text, "Chapter Heading")


def _add_section_heading(doc: Document, text: str):
    """節標題 → Section Heading 樣式。"""
    return _add_para(doc, text, "Section Heading")


def _add_subsection_heading(doc: Document, text: str):
    """子節標題 → SubSection Heading 樣式。"""
    return _add_para(doc, text, "SubSection Heading")


def _add_body(doc: Document, text: str):
    """內文 → Body Text 樣式。"""
    return _add_para(doc, text, "Body Text")


def _add_figure_caption(doc: Document, caption: str):
    """圖說（圖片下方）→ Figure Caption 樣式。"""
    return _add_para(doc, caption, "Figure Caption")


def _add_table_caption(doc: Document, caption: str):
    """表說（表格上方）→ Table Caption 樣式。"""
    return _add_para(doc, caption, "Table Caption")


def _add_page_break(doc: Document):
    """插入分頁符號（在 run 內以 XML 插入）。"""
    p = doc.add_paragraph(style=doc.styles["Normal"])
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ── 模板 Section 結構處理 ─────────────────────────────────────────────────────

def _extract_and_clear_body(doc: Document):
    """
    1. 從 body 中找到 embedded sectPr（定義前置區段格式）並深度複製保存。
    2. 清空 body 所有 <w:p> 和 <w:tbl>。
    3. 確保 body 最後有一個空 <w:p>（Word 規格要求）。

    回傳保存的 embedded sectPr（若模板未包含則回傳 None）。
    """
    body = doc.element.body
    saved_sect_pr = None

    # 找 embedded sectPr（在某個 <w:p> 的 <w:pPr> 底下）
    for child in body:
        if child.tag.split("}")[-1] == "p":
            pPr = child.find(qn("w:pPr"))
            if pPr is not None:
                sp = pPr.find(qn("w:sectPr"))
                if sp is not None:
                    saved_sect_pr = deepcopy(sp)
                    break

    # 移除所有段落與表格
    to_remove = [
        child for child in body
        if child.tag.split("}")[-1] in ("p", "tbl", "sdt")
    ]
    for elem in to_remove:
        body.remove(elem)

    # 確保 body 有一個空段落（Word 規格）
    placeholder = OxmlElement("w:p")
    body.append(placeholder)

    return saved_sect_pr


def _insert_section_boundary(doc: Document, sect_pr_template) -> None:
    """
    在目前文件末端插入一個攜帶 embedded sectPr 的空段落，
    作為前置區段（羅馬數字）與正文區段（阿拉伯數字）的分界點。

    sect_pr_template : 從模板保存的 sectPr（已含 pgNumType、pgSz、pgMar、footerReference…）
    """
    body = doc.element.body

    # 建立 <w:p><w:pPr><w:sectPr>…</w:sectPr></w:pPr></w:p>
    new_p  = OxmlElement("w:p")
    pPr    = OxmlElement("w:pPr")
    new_p.append(pPr)

    if sect_pr_template is not None:
        pPr.append(deepcopy(sect_pr_template))
    else:
        # 後備：建立最基本的 nextPage 分節，前置用 lowerRoman
        sp = OxmlElement("w:sectPr")
        w_type = OxmlElement("w:type")
        w_type.set(qn("w:val"), "nextPage")
        sp.append(w_type)
        pg_num = OxmlElement("w:pgNumType")
        pg_num.set(qn("w:fmt"), "lowerRoman")
        pg_num.set(qn("w:start"), "1")
        sp.append(pg_num)
        pPr.append(sp)

    # 插入到 body 末端（在 body-level sectPr 之前）
    body_sect = body.find(qn("w:sectPr"))
    if body_sect is not None:
        body_sect.addprevious(new_p)
    else:
        body.append(new_p)


# ── 各區段建構 ────────────────────────────────────────────────────────────────

def _build_cover(doc: Document, cover: CoverInfo, school_name: str | None) -> None:
    """封面頁。全部使用 Normal + 手動對齊，避免干擾前置樣式。"""
    univ   = cover.university or school_name or ""
    dept   = cover.department or ""
    degree = cover.degree or "碩士"

    if univ:
        _add_para(doc, univ,   "Normal", bold=True,  font_size_pt=18,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if dept:
        _add_para(doc, dept,   "Normal", bold=True,  font_size_pt=16,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, f"{degree}論文", "Normal", font_size_pt=14,
              alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    if cover.title_zh:
        _add_para(doc, cover.title_zh, "Normal", bold=True, font_size_pt=20,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if cover.title_en:
        _add_para(doc, cover.title_en, "Normal", font_size_pt=14,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    if cover.author:
        _add_para(doc, f"研究生：{cover.author}",   "Normal",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if cover.student_id:
        _add_para(doc, f"學　號：{cover.student_id}", "Normal",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if cover.advisor:
        _add_para(doc, f"指導教授：{cover.advisor} 博士", "Normal",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if cover.co_advisor:
        _add_para(doc, f"共同指導：{cover.co_advisor} 博士", "Normal",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    year  = cover.year  or ""
    month = cover.month or ""
    if year or month:
        _add_para(doc, f"中華民國 {year} 年 {month} 月", "Normal",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    _add_page_break(doc)


def _build_abstract(doc: Document, abstract: AbstractSection, label: str) -> None:
    """摘要頁（中文 / 英文）。"""
    _add_front_matter_title(doc, label)
    _add_body(doc, abstract.content)
    if abstract.keywords:
        sep    = "、" if "中文" in label else ", "
        prefix = "關鍵字：" if "中文" in label else "Keywords: "
        _add_body(doc, f"{prefix}{sep.join(abstract.keywords)}")
    _add_page_break(doc)


def _build_acknowledgments(doc: Document, text: str) -> None:
    """誌謝頁。"""
    _add_front_matter_title(doc, "誌　謝")
    _add_body(doc, text)
    _add_page_break(doc)


def _build_toc_placeholder(doc: Document) -> None:
    """目錄（placeholder，需在 Word 中手動更新）。"""
    _add_front_matter_title(doc, "目　錄")
    _add_para(
        doc,
        "（請在 Word 中點選此處，執行「更新功能變數」→「更新整個目錄」）",
        "Normal",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_page_break(doc)


def _render_blocks(doc: Document, blocks: list[ContentBlock]) -> None:
    """渲染 ContentBlock 清單至文件。"""
    for block in blocks:
        if isinstance(block, ParagraphBlock):
            _add_body(doc, block.text)
        elif isinstance(block, FigureBlock):
            # 圖片無法從 docx 萃取，以 placeholder 框代替
            _add_para(
                doc,
                f"【圖片】{block.caption}",
                "Normal",
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _add_figure_caption(doc, block.caption)   # 圖說在下方
        elif isinstance(block, TableBlock):
            _add_table_block(doc, block)


def _add_table_block(doc: Document, block: TableBlock) -> None:
    """表說（上方）+ 表格資料。"""
    _add_table_caption(doc, block.caption)   # 表說在上方

    if not block.rows:
        _add_body(doc, "（表格資料暫缺）")
        return

    col_count = max(len(r) for r in block.rows)
    tbl = doc.add_table(rows=len(block.rows), cols=col_count)
    try:
        tbl.style = doc.styles["Table Grid"]
    except KeyError:
        pass
    for r_idx, row_data in enumerate(block.rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < col_count:
                tbl.rows[r_idx].cells[c_idx].text = cell_text


def _build_chapter(doc: Document, chapter: Chapter) -> None:
    """一個章（含節 / 子節）。"""
    zh = _zh_num(chapter.number)
    _add_chapter_heading(doc, f"第{zh}章　{chapter.title}")
    _render_blocks(doc, chapter.intro_content)
    for sec in chapter.sections:
        _build_section(doc, sec)
    _add_page_break(doc)


def _build_section(doc: Document, section: Section) -> None:
    _add_section_heading(doc, f"{section.number}　{section.title}")
    _render_blocks(doc, section.content)
    for sub in section.subsections:
        _build_subsection(doc, sub)


def _build_subsection(doc: Document, sub: SubSection) -> None:
    _add_subsection_heading(doc, f"{sub.number}　{sub.title}")
    _render_blocks(doc, sub.content)


def _build_references(doc: Document, references: list[str]) -> None:
    """參考文獻。"""
    _add_chapter_heading(doc, "參考文獻")
    for ref in references:
        _add_body(doc, ref)


# ── 主要入口 ──────────────────────────────────────────────────────────────────

def build(
    schema: PaperSchema,
    template_path: "Path | str | None",
    config: "dict | None" = None,
) -> Document:
    """
    根據 PaperSchema 與 Word 模板產生格式化論文文件。

    schema        : ai_analyzer.analyze() 回傳的論文結構
    template_path : templates/ncu.docx 路徑（None 時建立無樣式空白文件）
    config        : 學校 JSON config（用於頁邊距覆寫；模板已有預設值）

    回傳           : docx.Document
    """
    # ── 1. 開啟模板 ───────────────────────────────────────────────────────────
    if template_path and Path(str(template_path)).exists():
        doc = Document(str(template_path))
    else:
        doc = Document()

    # ── 2. 清空 body，保存 embedded sectPr ───────────────────────────────────
    saved_sect_pr = _extract_and_clear_body(doc)

    # ── 3. 覆寫頁邊距（若 config 有指定）────────────────────────────────────
    if config:
        m = config.get("margins", {})
        for sec in doc.sections:
            if m.get("top_cm"):    sec.top_margin    = Cm(m["top_cm"])
            if m.get("bottom_cm"): sec.bottom_margin = Cm(m["bottom_cm"])
            if m.get("left_cm"):   sec.left_margin   = Cm(m["left_cm"])
            if m.get("right_cm"):  sec.right_margin  = Cm(m["right_cm"])

    # ── 4. 前置內容（前置區段：羅馬數字頁碼）────────────────────────────────
    school_name = config.get("school_name") if config else None

    _build_cover(doc, schema.cover, school_name)

    if schema.abstract_zh:
        _build_abstract(doc, schema.abstract_zh, "摘　要")
    if schema.abstract_en:
        _build_abstract(doc, schema.abstract_en, "ABSTRACT")
    if schema.acknowledgments:
        _build_acknowledgments(doc, schema.acknowledgments)

    _build_toc_placeholder(doc)

    # ── 5. 分節點：前置 → 正文 ───────────────────────────────────────────────
    _insert_section_boundary(doc, saved_sect_pr)

    # ── 6. 正文各章（正文區段：阿拉伯數字頁碼）──────────────────────────────
    for chapter in schema.chapters:
        _build_chapter(doc, chapter)

    # ── 7. 參考文獻 ───────────────────────────────────────────────────────────
    if schema.references:
        _build_references(doc, schema.references)

    return doc
