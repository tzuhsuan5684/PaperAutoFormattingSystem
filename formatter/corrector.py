"""
論文格式自動修正模組。
依照學校 JSON config 修正 .docx 文件的頁邊距、行距、頁碼、圖表說明位置等。
"""
import re
from copy import deepcopy
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.enum.text import WD_LINE_SPACING


# ── 正則：圖說 / 表說 / 資料來源 ──────────────────────────────────────────
_RE_FIGURE = re.compile(r"^(圖|Figure)\s*\d+", re.IGNORECASE)
_RE_TABLE  = re.compile(r"^(表|Table)\s*\d+",  re.IGNORECASE)
_RE_SOURCE = re.compile(r"^(資料來源|Source\s*[:：])", re.IGNORECASE)


class ThesisCorrector:
    """依照 config 自動修正論文格式。"""

    # ── 1. 頁邊距 ──────────────────────────────────────────────────────────
    def fix_margins(self, doc: Document, config: dict[str, Any]) -> None:
        """套用頁邊距到所有 section。"""
        m = config["margins"]
        top    = Cm(m["top_cm"])
        bottom = Cm(m["bottom_cm"])
        left   = Cm(m["left_cm"])
        right  = Cm(m["right_cm"])
        print(f"top: {top}, bottom: {bottom}, left: {left}, right: {right}")

        for section in doc.sections:
            section.top_margin    = top
            section.bottom_margin = bottom
            section.left_margin   = left
            section.right_margin  = right

    # ── 2. 行距 ────────────────────────────────────────────────────────────
    def fix_line_spacing(self, doc: Document, config: dict[str, Any]) -> None:
        """所有段落套用 1.5 倍行距；章標題後套用雙倍行距。"""
        ls_cfg  = config["line_spacing"]
        spacing = ls_cfg["value"]  # 1.5

        for para in doc.paragraphs:
            pf = para.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing      = spacing

            # 章標題段落後加大行距（雙倍）
            style_name = (para.style.name or "").lower()
            if "heading 1" in style_name or "chapter" in style_name:
                pf.space_after = Pt(24)

    # ── 3. 頁碼 ────────────────────────────────────────────────────────────
    def fix_page_numbers(self, doc: Document, config: dict[str, Any]) -> None:
        """
        在文件中插入頁碼：
        - 前置部分（第一個 section）：小寫羅馬數字，置中底端
        - 正文部分（第二個 section 起，或在遇到「第一章」時建立新 section）：
          阿拉伯數字從 1 開始，置中底端

        python-docx 的 section break 需要操作底層 XML。
        """
        pg_cfg = config["page_number"]

        # 確保至少有兩個 section（前置 / 正文）
        self._ensure_two_sections(doc)

        sections = list(doc.sections)

        # 前置：section 0
        self._set_footer_page_number(
            sections[0],
            fmt="lowerRoman",
            start=1,
            distance_cm=pg_cfg["distance_from_bottom_cm"],
        )

        # 正文：section 1+
        for sec in sections[1:]:
            self._set_footer_page_number(
                sec,
                fmt="decimal",
                start=1,
                distance_cm=pg_cfg["distance_from_bottom_cm"],
            )

    def _ensure_two_sections(self, doc: Document) -> None:
        """若文件只有一個 section，在偵測到正文第一章時插入分節符號。

        插入時從原有的 body-level sectPr 複製 pgSz / pgMar，確保新 Section[0]
        具備完整的頁面屬性，讓 docx-preview 能正確渲染封面頁。
        """
        if len(doc.sections) >= 2:
            return

        # 記錄原始 body sectPr，以便複製 pgSz / pgMar 到新 Section[0]
        body_sect_pr = doc.sections[0]._sectPr

        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower()
            text       = para.text.strip()
            is_chapter = (
                "heading 1" in style_name
                or re.match(r"^第[一二三四五六七八九十]+章", text)
            )
            if is_chapter:
                self._insert_section_break_before(
                    para,
                    break_type="nextPage",
                    ref_sect_pr=body_sect_pr,
                )
                break

    @staticmethod
    def _insert_section_break_before(
        para,
        break_type: str = "nextPage",
        ref_sect_pr=None,
    ) -> None:
        """在指定段落之前插入分節符號（操作底層 XML）。

        ref_sect_pr: 若提供，則從此 sectPr 複製 pgSz 與 pgMar 到新 Section，
                     確保封面頁有完整的頁面尺寸定義。
        """
        from copy import deepcopy

        new_p  = OxmlElement("w:p")
        pPr    = OxmlElement("w:pPr")
        sectPr = OxmlElement("w:sectPr")

        # 分節類型
        w_type = OxmlElement("w:type")
        w_type.set(qn("w:val"), break_type)
        sectPr.append(w_type)

        # 從參考 sectPr 複製頁面尺寸與邊距，避免封面頁 section 屬性不完整
        if ref_sect_pr is not None:
            for tag_name in ("w:pgSz", "w:pgMar"):
                elem = ref_sect_pr.find(qn(tag_name))
                if elem is not None:
                    sectPr.append(deepcopy(elem))

        pPr.append(sectPr)
        new_p.append(pPr)
        para._p.addprevious(new_p)

    @staticmethod
    def _set_footer_page_number(
        section,
        fmt: str,
        start: int,
        distance_cm: float,
    ) -> None:
        """
        在 section 的 footer 中心插入頁碼欄位，設定格式與起始頁。

        fmt 可為 "decimal" | "lowerRoman" | "upperRoman" 等（OOXML numFmt 值）。
        """
        # 設定 footer 距底端距離
        section.footer_distance = Cm(distance_cm)
        section.different_first_page_header_footer = False

        footer = section.footer
        # 清空現有 footer 段落
        for p in footer.paragraphs:
            p.clear()

        # 取得（或建立）第一個段落
        if footer.paragraphs:
            fp = footer.paragraphs[0]
        else:
            fp = footer.add_paragraph()

        # 置中
        fp.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

        # 插入頁碼欄位 XML： <w:fldChar> + <w:instrText> + <w:fldChar>
        run = fp.add_run()
        r_elem = run._r

        def _make_fld(fld_type: str):
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), fld_type)
            return fc

        r_elem.append(_make_fld("begin"))

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        r_elem.append(instr)
        r_elem.append(_make_fld("separate"))
        r_elem.append(_make_fld("end"))

        # 設定頁碼格式與起始頁（操作 sectPr）
        sect_pr = section._sectPr
        # 移除舊的 pgNumType
        for old in sect_pr.findall(qn("w:pgNumType")):
            sect_pr.remove(old)
        pg_num_type = OxmlElement("w:pgNumType")
        pg_num_type.set(qn("w:fmt"),   fmt)
        pg_num_type.set(qn("w:start"), str(start))
        sect_pr.append(pg_num_type)

    # ── 4. 圖表說明位置 ────────────────────────────────────────────────────
    def fix_figure_table_captions(self, doc: Document, _config: dict[str, Any] = None) -> None:
        """
        偵測圖說 / 表說 / 資料來源，確保：
        - 圖說在圖片（含圖片的段落）下方
        - 表說在表格上方
        - 資料來源永遠在圖表下方
        """
        body_elem = doc.element.body
        children  = list(body_elem)  # w:p 或 w:tbl

        i = 0
        while i < len(children):
            child = children[i]
            tag   = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            # ── 圖說：應在含圖片的段落後面 ──
            if tag == "p":
                text = _para_text(child)
                if _RE_FIGURE.match(text):
                    # 往前找圖片段落
                    if i > 0:
                        prev = children[i - 1]
                        prev_tag = prev.tag.split("}")[-1] if "}" in prev.tag else prev.tag
                        # 若前一個不是含圖片的段落，把圖說移到圖片下方
                        # （MVP：已在圖片後則不動；若偵測到圖說在圖片前則交換）
                        if prev_tag == "p" and not _has_image(prev):
                            # 找最近的含圖片段落
                            img_idx = self._find_nearest_image_before(children, i)
                            if img_idx is not None and img_idx != i - 1:
                                # 移動圖說到圖片下方
                                body_elem.remove(child)
                                children.pop(i)
                                insert_at = img_idx + 1
                                ref_elem  = children[insert_at] if insert_at < len(children) else None
                                if ref_elem is not None:
                                    body_elem.insert(list(body_elem).index(ref_elem), child)
                                else:
                                    body_elem.append(child)
                                children.insert(insert_at, child)
                                i = insert_at + 1
                                continue

            # ── 表說：應在 <w:tbl> 上方 ──
            if tag == "p":
                text = _para_text(child)
                if _RE_TABLE.match(text):
                    if i + 1 < len(children):
                        nxt     = children[i + 1]
                        nxt_tag = nxt.tag.split("}")[-1] if "}" in nxt.tag else nxt.tag
                        if nxt_tag == "tbl":
                            pass  # 已在正確位置
                        else:
                            # 找後方最近的表格
                            tbl_idx = self._find_nearest_table_after(children, i)
                            if tbl_idx is not None:
                                body_elem.remove(child)
                                children.pop(i)
                                tbl_idx -= 1  # 因為移除後 index 前移
                                ref_elem = children[tbl_idx]
                                body_elem.insert(list(body_elem).index(ref_elem), child)
                                children.insert(tbl_idx, child)
                                i = tbl_idx + 1
                                continue

            # ── 資料來源：移到最近圖/表後方 ──
            if tag == "p":
                text = _para_text(child)
                if _RE_SOURCE.match(text):
                    # 確保資料來源不在圖/表之前
                    if i + 1 < len(children):
                        nxt_tag = children[i + 1].tag.split("}")[-1]
                        if nxt_tag in ("drawing", "tbl"):
                            # 資料來源跑到圖表前面了，移到後面
                            target = children[i + 1]
                            body_elem.remove(child)
                            children.pop(i)
                            insert_at = children.index(target) + 1
                            ref_elem  = children[insert_at] if insert_at < len(children) else None
                            if ref_elem is not None:
                                body_elem.insert(list(body_elem).index(ref_elem), child)
                            else:
                                body_elem.append(child)
                            children.insert(insert_at, child)
                            i = insert_at + 1
                            continue

            i += 1

    @staticmethod
    def _find_nearest_image_before(children, current_idx: int):
        for j in range(current_idx - 1, -1, -1):
            if _has_image(children[j]):
                return j
        return None

    @staticmethod
    def _find_nearest_table_after(children, current_idx: int):
        for j in range(current_idx + 1, len(children)):
            tag = children[j].tag.split("}")[-1] if "}" in children[j].tag else children[j].tag
            if tag == "tbl":
                return j
        return None

    # ── 5. 字型修正 ────────────────────────────────────────────────────────
    def fix_fonts(self, doc: Document, config: dict[str, Any]) -> None:
        """
        逐一修正所有段落的每個 run：
          - 中文字型（w:eastAsia）→ config font.main_chinese（預設 標楷體）
          - 英文字型（w:ascii / w:hAnsi / w:cs）→ config font.main_english（預設 Times New Roman）
          - 字型大小 → config font.size_pt（預設 12pt）

        例外：樣式名稱含 "Heading" / "heading" / "ChapterHeading" /
              "SectionHeading" / "SubSectionHeading" 的段落，
              只修字型名稱，不改字型大小（保留標題的大字號）。
        """
        font_cfg = config.get("font", {})
        zh_font  = font_cfg.get("main_chinese", "標楷體")
        en_font  = font_cfg.get("main_english", "Times New Roman")
        size_pt  = font_cfg.get("size_pt", 12)
        half_pt  = str(size_pt * 2)   # OOXML 用 half-points

        _HEADING_KEYWORDS = ("heading", "chapterheading", "sectionheading", "subsectionheading")

        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower().replace(" ", "")
            is_heading  = any(kw in style_name for kw in _HEADING_KEYWORDS)

            for run in para.runs:
                r_elem = run._r

                # 取得或建立 rPr
                rpr = r_elem.find(qn("w:rPr"))
                if rpr is None:
                    rpr = OxmlElement("w:rPr")
                    r_elem.insert(0, rpr)

                # ── 字型名稱 ──
                rFonts = rpr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rpr.insert(0, rFonts)
                rFonts.set(qn("w:ascii"),    en_font)
                rFonts.set(qn("w:hAnsi"),    en_font)
                rFonts.set(qn("w:eastAsia"), zh_font)
                rFonts.set(qn("w:cs"),       en_font)

                # ── 字型大小（標題段落跳過）──
                if not is_heading:
                    for tag in ("w:sz", "w:szCs"):
                        sz_elem = rpr.find(qn(tag))
                        if sz_elem is None:
                            sz_elem = OxmlElement(tag)
                            rpr.append(sz_elem)
                        sz_elem.set(qn("w:val"), half_pt)

    # ── 6. run_all ─────────────────────────────────────────────────────────
    def run_all(
        self,
        doc: Document,
        config: dict[str, Any],
        template_path=None,
    ) -> Document:
        """依序執行所有格式修正，回傳修正後的 doc。

        執行順序說明：
        0. （選填）inject_template_styles：從學校模板注入 Word Styles
        1. fix_page_numbers 先建立前置／正文兩個 section 結構（含 pgSz 複製）
        2. fix_margins 再對所有 section 套用邊距，確保新建的 Section[0] 也被涵蓋
        3. fix_line_spacing
        4. fix_fonts：字型名稱 + 字型大小
        5. fix_figure_table_captions

        template_path: Path | None
            傳入 templates/{school_id}.docx 的路徑；
            None 時跳過樣式注入，退回純 JSON config 修正模式。
        """
        if template_path is not None:
            self._inject_template_styles(doc, template_path)
        self.fix_page_numbers(doc, config)
        self.fix_margins(doc, config)
        self.fix_line_spacing(doc, config)
        self.fix_fonts(doc, config)
        self.fix_figure_table_captions(doc, config)
        return doc

    @staticmethod
    def _inject_template_styles(doc: Document, template_path) -> None:
        """將模板的 Word Styles 注入 doc（已有的不覆蓋）。"""
        from formatter.template_loader import inject_styles_from_template
        injected = inject_styles_from_template(template_path, doc)
        if injected:
            print(f"[template] 注入 {len(injected)} 個樣式：{injected}")


# ── 工具函式 ────────────────────────────────────────────────────────────────

def _para_text(p_elem) -> str:
    """取得 <w:p> 元素的純文字。"""
    return "".join(t.text or "" for t in p_elem.iter(qn("w:t")))


def _has_image(p_elem) -> bool:
    """判斷 <w:p> 元素中是否含有圖片（drawing / pict）。"""
    for tag in (qn("w:drawing"), qn("w:pict"), qn("a:blip")):
        if p_elem.find(".//" + tag) is not None:
            return True
    return False
