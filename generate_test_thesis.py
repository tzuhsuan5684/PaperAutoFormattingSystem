"""
產生測試用假論文 test_thesis.docx
故意包含一些格式錯誤，方便測試修正功能。
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── 封面 ────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("國立中央大學")
run.font.size = Pt(18)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("資訊工程學系")
run.font.size = Pt(16)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("碩士論文")
run.font.size = Pt(20)
run.font.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("基於深度學習之自然語言處理\n應用於論文格式自動檢查之研究")
run.font.size = Pt(16)
run.font.bold = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("研究生：王小明\n指導教授：李大偉 博士\n\n中華民國一一三年六月")

doc.add_page_break()

# ── 書名頁 ───────────────────────────────────────────
p = doc.add_heading("書名頁", level=1)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("（此頁為書名頁，內容同封面）")
doc.add_page_break()

# ── 中文摘要 ─────────────────────────────────────────
doc.add_heading("中文摘要", level=1)
doc.add_paragraph(
    "隨著人工智慧技術的快速發展，自然語言處理已被廣泛應用於各類文件處理任務中。"
    "本研究提出一套基於深度學習的論文格式自動檢查系統，透過訓練專用的語言模型，"
    "能夠有效識別論文中的格式錯誤，包含字型、行距、頁邊距、圖表說明位置等問題。"
    "實驗結果顯示，本系統在測試集上達到 92.3% 的格式問題偵測準確率，"
    "相較於傳統規則式方法提升了 15.7 個百分點。"
    "此外，系統亦提供自動修正功能，可直接將論文修正至符合學校規範的格式，"
    "大幅節省研究生花費在格式調整上的時間，具有相當的實用價值。"
)
doc.add_paragraph(
    "關鍵字：深度學習、自然語言處理、論文格式檢查、自動修正、文件處理"
)
doc.add_page_break()

# ── 英文摘要 ─────────────────────────────────────────
doc.add_heading("Abstract", level=1)
doc.add_paragraph(
    "With the rapid advancement of artificial intelligence, natural language processing "
    "has been widely applied to various document processing tasks. This study proposes "
    "an automatic thesis format checking system based on deep learning. By training a "
    "dedicated language model, the system can effectively identify formatting errors in "
    "theses, including font settings, line spacing, page margins, and figure caption positions. "
    "Experimental results show that the system achieves a format error detection accuracy "
    "of 92.3% on the test set, improving 15.7 percentage points over traditional rule-based methods. "
    "Furthermore, the system provides automatic correction functionality that directly revises "
    "the thesis to conform to school formatting standards, significantly reducing the time "
    "graduate students spend on format adjustments."
)
doc.add_paragraph(
    "Keywords: deep learning, natural language processing, thesis format checking, "
    "automatic correction, document processing"
)
doc.add_page_break()

# ── 誌謝 ─────────────────────────────────────────────
doc.add_heading("誌謝", level=1)
doc.add_paragraph(
    "首先，衷心感謝指導教授李大偉博士在研究過程中給予的耐心指導與寶貴建議。"
    "教授嚴謹的治學態度與豐富的學術經驗，讓我獲益匪淺，也讓本研究得以順利完成。"
)
doc.add_paragraph(
    "同時感謝口試委員張三教授、陳四教授提供的寶貴意見，使本論文更加完善。"
    "也感謝實驗室的同學們，在研究過程中互相扶持、討論切磋，留下了許多美好的回憶。"
)
doc.add_paragraph("最後，謹將此論文獻給我的家人，感謝他們一路以來的支持與鼓勵。")
doc.add_page_break()

# ── 目錄 ─────────────────────────────────────────────
doc.add_heading("目錄", level=1)
toc_items = [
    ("中文摘要", "i"),
    ("Abstract", "ii"),
    ("誌謝", "iii"),
    ("目錄", "iv"),
    ("圖目錄", "v"),
    ("表目錄", "vi"),
    ("第一章　緒論", "1"),
    ("　1-1　研究背景與動機", "1"),
    ("　1-2　研究目的", "2"),
    ("　1-3　論文架構", "3"),
    ("第二章　相關研究", "4"),
    ("　2-1　深度學習基礎", "4"),
    ("　2-2　自然語言處理", "5"),
    ("第三章　系統設計", "7"),
    ("　3-1　系統架構", "7"),
    ("　3-2　模型訓練", "9"),
    ("第四章　實驗結果", "11"),
    ("第五章　結論與未來展望", "14"),
    ("參考文獻", "15"),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(item)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14))
    p.add_run(f"\t{page}")
doc.add_page_break()

# ── 圖目錄 ───────────────────────────────────────────
doc.add_heading("圖目錄", level=1)
fig_items = [
    ("圖 1-1　論文格式自動修正系統示意圖", "3"),
    ("圖 3-1　系統整體架構圖", "7"),
    ("圖 3-2　模型訓練流程圖", "9"),
    ("圖 4-1　各模型準確率比較", "12"),
]
for item, page in fig_items:
    p = doc.add_paragraph()
    p.add_run(item)
doc.add_page_break()

# ── 表目錄 ───────────────────────────────────────────
doc.add_heading("表目錄", level=1)
tbl_items = [
    ("表 3-1　資料集統計資訊", "8"),
    ("表 4-1　實驗結果比較", "11"),
    ("表 4-2　各類格式錯誤偵測率", "13"),
]
for item, page in tbl_items:
    p = doc.add_paragraph()
    p.add_run(item)
doc.add_page_break()

# ── 第一章 ───────────────────────────────────────────
doc.add_heading("第一章　緒論", level=1)

doc.add_heading("1-1　研究背景與動機", level=2)
doc.add_paragraph(
    "論文撰寫是研究生學術生涯中不可或缺的重要環節。然而，各大學對於論文格式均有"
    "嚴格且細緻的規範，包含頁邊距、字型、行距、章節編號格式、圖表說明位置等諸多項目。"
    "根據調查，超過八成的研究生表示曾因格式問題而導致論文審核延誤，"
    "平均每人花費在格式調整上的時間超過十小時。"
)
doc.add_paragraph(
    "有鑑於此，本研究希望藉由人工智慧技術，開發一套能夠自動檢查並修正論文格式的系統，"
    "協助研究生快速完成格式調整，將更多精力投入研究內容本身。"
)

doc.add_heading("1-2　研究目的", level=2)
doc.add_paragraph("本研究的主要目的如下：")
doc.add_paragraph("（1）建立一套能夠自動識別論文格式錯誤的深度學習模型。", style='List Number')
doc.add_paragraph("（2）實現論文格式的自動修正功能，支援主流學校格式規範。", style='List Number')
doc.add_paragraph("（3）提供使用者友善的操作介面，降低使用門檻。", style='List Number')

doc.add_heading("1-3　論文架構", level=2)
doc.add_paragraph(
    "本論文共分為五章。第一章為緒論，說明研究背景、動機與目的。"
    "第二章回顧相關研究，包含深度學習與自然語言處理的相關文獻。"
    "第三章說明系統設計與模型訓練方法。"
    "第四章呈現實驗結果與分析。"
    "第五章為結論與未來展望。"
)
doc.add_page_break()

# ── 第二章 ───────────────────────────────────────────
doc.add_heading("第二章　相關研究", level=1)

doc.add_heading("2-1　深度學習基礎", level=2)
doc.add_paragraph(
    "深度學習（Deep Learning）是機器學習的一個子領域，透過多層神經網路"
    "學習資料的抽象表徵。自 2012 年 AlexNet 在 ImageNet 競賽中取得突破性成果後，"
    "深度學習已成為人工智慧領域最重要的技術之一 [1]。"
)
doc.add_paragraph(
    "Transformer 架構 [2] 的提出更是帶來了自然語言處理領域的革命性突破，"
    "後續發展出的 BERT [3]、GPT [4] 等預訓練模型大幅提升了各類 NLP 任務的表現。"
)

doc.add_heading("2-2　自然語言處理", level=2)
doc.add_paragraph(
    "自然語言處理（Natural Language Processing, NLP）是人工智慧與語言學的交叉領域，"
    "旨在使電腦能夠理解、生成和處理人類語言。近年來，基於大型語言模型的方法"
    "在文字分類、命名實體識別、機器翻譯等任務上均取得了顯著進展。"
)

# 故意把圖說放在圖片前（製造格式錯誤）
p = doc.add_paragraph()
p.add_run("圖 2-1　Transformer 模型架構示意圖")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("（此處應有圖片）")
doc.add_page_break()

# ── 第三章 ───────────────────────────────────────────
doc.add_heading("第三章　系統設計", level=1)

doc.add_heading("3-1　系統架構", level=2)
doc.add_paragraph(
    "本系統採用前後端分離的架構設計。後端以 Python FastAPI 框架實作 RESTful API，"
    "前端則以現代化的 Web 技術開發使用者介面。系統主要分為三個模組："
    "格式檢查模組（Checker）、格式修正模組（Corrector）與文件匯出模組（Exporter）。"
)

# 表格
doc.add_paragraph("表 3-1　資料集統計資訊")
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
headers = ["資料集", "論文數量", "格式錯誤數"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ("訓練集", "1,200", "8,543"),
    ("驗證集", "200",   "1,421"),
    ("測試集", "200",   "1,389"),
]
for r, row in enumerate(data, 1):
    for c, val in enumerate(row):
        table.rows[r].cells[c].text = val

doc.add_paragraph()

doc.add_heading("3-2　模型訓練", level=2)
doc.add_paragraph(
    "本研究使用 BERT-base-chinese 作為預訓練模型的基礎，"
    "在自行標注的論文格式錯誤資料集上進行微調（fine-tuning）。"
    "訓練時使用 AdamW 優化器，學習率設為 2e-5，批次大小為 32，訓練 10 個 epoch。"
)
doc.add_page_break()

# ── 第四章 ───────────────────────────────────────────
doc.add_heading("第四章　實驗結果", level=1)

doc.add_paragraph(
    "本章節呈現系統在測試集上的實驗結果，並與現有方法進行比較。"
)

doc.add_paragraph("表 4-1　實驗結果比較")
table2 = doc.add_table(rows=4, cols=4)
table2.style = 'Table Grid'
h2 = ["方法", "Precision", "Recall", "F1-Score"]
for i, h in enumerate(h2):
    table2.rows[0].cells[i].text = h
d2 = [
    ("規則式方法",   "0.812", "0.743", "0.776"),
    ("傳統機器學習", "0.856", "0.821", "0.838"),
    ("本研究方法",   "0.931", "0.916", "0.923"),
]
for r, row in enumerate(d2, 1):
    for c, val in enumerate(row):
        table2.rows[r].cells[c].text = val

doc.add_paragraph()
doc.add_paragraph(
    "由表 4-1 可見，本研究方法在所有指標上均優於比較方法，"
    "F1-Score 達到 0.923，較規則式方法提升了 14.7 個百分點。"
)
doc.add_paragraph(
    "表 4-2 進一步分析各類格式錯誤的偵測率，其中頁邊距錯誤（98.2%）與"
    "行距錯誤（96.5%）的偵測率最高，圖說位置錯誤（87.3%）相對較低，"
    "主要原因為部分圖說位置的判斷需要理解文件語意脈絡。"
)
doc.add_page_break()

# ── 第五章 ───────────────────────────────────────────
doc.add_heading("第五章　結論與未來展望", level=1)

doc.add_paragraph(
    "本研究成功開發了一套基於深度學習的論文格式自動檢查與修正系統，"
    "在測試集上達到 92.3% 的格式問題偵測準確率。系統能夠有效識別並自動修正"
    "頁邊距、字型大小、行距、頁碼格式、圖表說明位置等常見格式問題，"
    "具有良好的實用性。"
)
doc.add_paragraph(
    "未來的研究方向包含：（1）擴充支援更多學校的格式規範；"
    "（2）加入語意理解能力，提升複雜格式錯誤的偵測率；"
    "（3）開發跨平台的桌面應用程式版本，提升使用便利性。"
)
doc.add_page_break()

# ── 參考文獻 ─────────────────────────────────────────
doc.add_heading("參考文獻", level=1)
refs = [
    "[1] Krizhevsky, A., Sutskever, I., & Hinton, G. E. (2012). ImageNet classification with deep convolutional neural networks. Advances in Neural Information Processing Systems, 25.",
    "[2] Vaswani, A., et al. (2017). Attention is all you need. Advances in Neural Information Processing Systems, 30.",
    "[3] Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2018). BERT: Pre-training of deep bidirectional transformers for language understanding. arXiv preprint arXiv:1810.04805.",
    "[4] Brown, T., et al. (2020). Language models are few-shot learners. Advances in Neural Information Processing Systems, 33.",
    "[5] 王小明、李大偉 (2024)。論文格式自動修正系統之研究。資訊工程學刊，12(3)，45-58。",
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.left_indent = Cm(0.6)

# ── 儲存 ─────────────────────────────────────────────
output = "test_thesis.docx"
doc.save(output)
print(f"Done: {output}")
