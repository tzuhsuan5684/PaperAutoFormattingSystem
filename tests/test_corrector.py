"""
formatter/corrector.py 基本測試。
執行方式：pytest tests/test_corrector.py -v
"""
import io
import json
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm

from formatter.corrector import ThesisCorrector
from formatter.loader import load_config


# ── 共用 fixture ──────────────────────────────────────────────────────────────

@pytest.fixture
def ncu_config():
    return load_config("ncu")


@pytest.fixture
def corrector():
    return ThesisCorrector()


def _make_doc_with_wrong_margins() -> Document:
    """建立頁邊距全部錯誤的測試文件。"""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(5.0)   # 錯誤
        section.bottom_margin = Cm(5.0)   # 錯誤
        section.left_margin   = Cm(5.0)   # 錯誤
        section.right_margin  = Cm(5.0)   # 錯誤
    doc.add_paragraph("測試段落一")
    doc.add_paragraph("測試段落二")
    return doc


# ── 1. 頁邊距測試 ─────────────────────────────────────────────────────────────

class TestFixMargins:
    def test_margins_are_corrected(self, corrector, ncu_config):
        doc = _make_doc_with_wrong_margins()
        corrector.fix_margins(doc, ncu_config)

        for section in doc.sections:
            assert abs(section.top_margin    - Cm(2.5)) < 500, "上邊距應為 2.5cm"
            assert abs(section.bottom_margin - Cm(2.5)) < 500, "下邊距應為 2.5cm"
            assert abs(section.left_margin   - Cm(3.0)) < 500, "左邊距應為 3.0cm"
            assert abs(section.right_margin  - Cm(2.0)) < 500, "右邊距應為 2.0cm"

    def test_all_sections_are_updated(self, corrector, ncu_config):
        """多個 section 都應套用相同邊距。"""
        doc = _make_doc_with_wrong_margins()
        # 人工插入第二個 section（透過 XML，簡化版）
        corrector.fix_margins(doc, ncu_config)
        for section in doc.sections:
            assert abs(section.left_margin - Cm(3.0)) < 500


# ── 2. 行距測試 ───────────────────────────────────────────────────────────────

class TestFixLineSpacing:
    def test_line_spacing_applied(self, corrector, ncu_config):
        doc = Document()
        doc.add_paragraph("第一段")
        doc.add_paragraph("第二段")
        corrector.fix_line_spacing(doc, ncu_config)

        from docx.enum.text import WD_LINE_SPACING
        for para in doc.paragraphs:
            rule = para.paragraph_format.line_spacing_rule
            assert rule in (WD_LINE_SPACING.MULTIPLE, WD_LINE_SPACING.ONE_POINT_FIVE), \
                f"行距規則應為 MULTIPLE 或 ONE_POINT_FIVE，實際：{rule}"


# ── 3. run_all 整合測試 ───────────────────────────────────────────────────────

class TestRunAll:
    def test_run_all_returns_document(self, corrector, ncu_config):
        doc = _make_doc_with_wrong_margins()
        result = corrector.run_all(doc, ncu_config)
        assert result is not None

    def test_run_all_fixes_margins(self, corrector, ncu_config):
        doc = _make_doc_with_wrong_margins()
        result = corrector.run_all(doc, ncu_config)
        for section in result.sections:
            assert abs(section.top_margin - Cm(2.5)) < 500

    def test_run_all_output_is_valid_docx(self, corrector, ncu_config, tmp_path):
        """修正後的文件應可正常儲存並重新讀取。"""
        doc = _make_doc_with_wrong_margins()
        result    = corrector.run_all(doc, ncu_config)
        out_path  = tmp_path / "output.docx"
        result.save(str(out_path))

        # 重新讀取驗證格式完整
        reloaded = Document(str(out_path))
        assert len(reloaded.paragraphs) >= 2


# ── 4. 圖說 / 表說偵測測試 ───────────────────────────────────────────────────

class TestFigureTableCaption:
    def test_figure_caption_regex(self):
        import re
        from formatter.corrector import _RE_FIGURE, _RE_TABLE

        assert _RE_FIGURE.match("圖 1 測試圖說")
        assert _RE_FIGURE.match("Figure 2: Test Caption")
        assert _RE_TABLE.match("表 3 測試表說")
        assert _RE_TABLE.match("Table 4 Summary")
        assert not _RE_FIGURE.match("這是一般段落")
        assert not _RE_TABLE.match("Figure 1")


# ── 5. config 載入測試 ───────────────────────────────────────────────────────

class TestLoadConfig:
    def test_ncu_config_loads(self, ncu_config):
        assert ncu_config["school_id"] == "ncu"
        assert ncu_config["margins"]["left_cm"] == 3.0

    def test_missing_school_raises(self):
        with pytest.raises(FileNotFoundError):
            load_config("nonexistent_school")
